import os
from decimal import Decimal
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user
import boto3
from botocore.exceptions import ClientError
from docx import Document
 

region_name = 'us-east-1'  
s3_bucket_name = 'jhdasfghjdbshbrwuerbhwehbgt'  
dynamodb = boto3.resource('dynamodb', region_name=region_name)
s3_client = boto3.client('s3', region_name=region_name)
sqs_client = boto3.client('sqs', region_name='us-east-1')  
queue_url = 'https://sqs.us-east-1.amazonaws.com/225489345950/BudgetLimitQueue'  
sns_topic_arn = 'arn:aws:sns:us-east-1:225489345950:BudgetLimitExceeded'
dynamodb = boto3.resource('dynamodb')
sns = boto3.client('sns')

sns.subscribe(
    TopicArn=sns_topic_arn,
    Protocol='email',  
    Endpoint='x23194383@student.ncirl.ie' 
)


table_expenses_name = 'your_expenses_table'

def fetch_all_expenses():
    table_expenses = dynamodb.Table(table_expenses_name)
    response = table_expenses.scan()  
    return response['Items']  

def check_budget_limit(expenses, budget_limit):
    try:
        total_expenses = sum(Decimal(expense['amount']) for expense in expenses)
        if total_expenses > budget_limit:
        
            message = f"Warning! Your expenses have exceeded the budget limit of ${budget_limit}. Total: ${total_expenses}."
            sns.publish(
                TopicArn=sns_topic_arn,
                Message=message,
                Subject='Budget Limit Exceeded'
            )
    except NoCredentialsError:
        print("Error: AWS credentials not found")
    except Exception as e:
        print(f"Error occurred while checking budget limit or sending SNS: {str(e)}")

def send_expense_to_sqs(expense_data):
    try:
        
        response = sqs_client.send_message(
            QueueUrl=queue_url,
            MessageBody=str(expense_data)  
        )
        print(f"Message sent to SQS with message ID: {response['MessageId']}")
    except Exception as e:
        print(f"Failed to send message to SQS: {e}")


table_expenses_name = 'Expenses'
table_income_name = 'Income'
table_budget_name = 'BudgetLimit'
table_users_name = 'Users'  


app = Flask(__name__)
app.secret_key = 'your_secret_key'  

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"  

class User(UserMixin):
    def __init__(self, id, username):
        self.id = id
        self.username = username


@login_manager.user_loader
def load_user(user_id):
    user_data = get_user_from_db(user_id)
    if user_data:
        return User(id=user_data['username'], username=user_data['username'])
    return None

def get_user_from_db(username):
    table_users = dynamodb.Table(table_users_name)
    response = table_users.get_item(Key={'username': username})
    return response.get('Item')

def create_dynamodb_tables():
    for table_name in [table_expenses_name, table_income_name, table_budget_name, table_users_name]:
        try:
            table = dynamodb.Table(table_name)
            table.load()
            print(f"DynamoDB Table '{table_name}' already exists.")
        except ClientError as e:
            if e.response['Error']['Code'] == 'ResourceNotFoundException':
                print(f"Creating DynamoDB Table: {table_name}")
                dynamodb.create_table(
                    TableName=table_name,
                    KeySchema=[{
                        'AttributeName': 'id' if table_name != table_users_name else 'username',
                        'KeyType': 'HASH'
                    }],
                    AttributeDefinitions=[{
                        'AttributeName': 'id' if table_name != table_users_name else 'username',
                        'AttributeType': 'S'
                    }],
                    ProvisionedThroughput={
                        'ReadCapacityUnits': 5,
                        'WriteCapacityUnits': 5
                    }
                )
                print(f"DynamoDB Table '{table_name}' created successfully.")
            else:
                print(f"Error checking DynamoDB table: {e}")
                raise

def create_s3_bucket():
    try:
       
        s3_client.head_bucket(Bucket=s3_bucket_name)
        print(f"S3 Bucket '{s3_bucket_name}' already exists.")
    except ClientError as e:
        if e.response['Error']['Code'] == 'NotFound':
            print(f"Creating S3 Bucket: {s3_bucket_name}")
            s3_client.create_bucket(Bucket=s3_bucket_name, CreateBucketConfiguration={'LocationConstraint': region_name})
            print(f"S3 Bucket '{s3_bucket_name}' created successfully.")
        else:
            print(f"Error checking or creating S3 bucket: {e}")
            raise


@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password'] 
        user_data = get_user_from_db(username)

        if user_data:
            flash('User already exists!', 'danger')
            return redirect(url_for('register'))
        else:
            table_users = dynamodb.Table(table_users_name)
            table_users.put_item(Item={'username': username, 'password': password})  
            flash('User registered successfully!', 'success')
            return redirect(url_for('login'))

    return render_template('register.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user_data = get_user_from_db(username)

        if user_data and user_data['password'] == password: 
            user = User(id=username, username=username)
            login_user(user)
            flash('Login successful!', 'success')
            return redirect(url_for('index'))
        else:
            flash('Invalid credentials!', 'danger')

    return render_template('login.html')


@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    try:
        table_expenses = dynamodb.Table(table_expenses_name)
        table_income = dynamodb.Table(table_income_name)
        table_budget = dynamodb.Table(table_budget_name)

       
        expense_response = table_expenses.scan()
        expenses = expense_response.get('Items', [])

        income_response = table_income.scan()
        incomes = income_response.get('Items', [])

       
        budget_response = table_budget.get_item(Key={'id': 'monthly_limit'})
        budget_limit = budget_response.get('Item', {}).get('limit', 0)

       
        total_expenses = sum(float(exp['amount']) for exp in expenses)
        total_income = sum(float(inc['amount']) for inc in incomes)
        balance = total_income - total_expenses

        
        budget_exceeded = total_expenses > budget_limit

        return render_template('dashboard.html', 
                               expenses=expenses, 
                               incomes=incomes, 
                               balance=balance, 
                               budget_limit=budget_limit, 
                               budget_exceeded=budget_exceeded, 
                               total_expenses=total_expenses, 
                               total_income=total_income)
    
    except ClientError as e:
        error_message = f"Error fetching data from DynamoDB: {e.response['Error']['Message']}"
        print(error_message)
        return error_message
   
@app.route('/edit_expense/<expense_id>', methods=['GET', 'POST'])
@login_required
def edit_expense(expense_id):
   
    table_expenses = dynamodb.Table(table_expenses_name)
    response = table_expenses.get_item(Key={'id': expense_id})
    expense = response.get('Item')
    
    if not expense:
        flash('Expense not found!', 'danger')
        return redirect(url_for('index'))
    
    if request.method == 'POST':
       
        expense_name = request.form['name']
        expense_amount = Decimal(request.form['amount'])
        expense_category = request.form['category']

       
        table_expenses.update_item(
            Key={'id': expense_id},
            UpdateExpression="set #name=:n, amount=:a, category=:c",
            ExpressionAttributeNames={"#name": "name"},
            ExpressionAttributeValues={":n": expense_name, ":a": expense_amount, ":c": expense_category}
        )

        flash('Expense updated successfully!', 'success')
        return redirect(url_for('index'))

   
    return render_template('edit_expense.html', expense=expense)

@app.route('/delete_expense/<expense_id>', methods=['GET'])
@login_required
def delete_expense(expense_id):
   
    table_expenses = dynamodb.Table(table_expenses_name)
    table_expenses.delete_item(Key={'id': expense_id})

    flash('Expense deleted successfully!', 'success')
    return redirect(url_for('index'))


@app.route('/edit_income/<income_id>', methods=['GET', 'POST'])
def edit_income(income_id):
   
    income_categories = ['Salary', 'Business', 'Investment', 'Other']

    table_income = dynamodb.Table(table_income_name)

    if request.method == 'POST':
        income_name = request.form['name']
        income_amount = Decimal(request.form['amount']) 
        income_category = request.form['category'] 

        try:
            table_income.update_item(
                Key={'id': income_id},
                UpdateExpression="SET #n = :name, #a = :amount, #c = :category",
                ExpressionAttributeNames={
                    '#n': 'name',
                    '#a': 'amount',
                    '#c': 'category'
                },
                ExpressionAttributeValues={
                    ':name': income_name,
                    ':amount': income_amount,
                    ':category': income_category
                }
            )
            flash('Income updated successfully!', 'success')
            return redirect(url_for('index'))
        except ClientError as e:
            error_message = f"Error updating income in DynamoDB: {e.response['Error']['Message']}"
            print(error_message)
            flash('Error updating income!', 'danger')
            return error_message

    try:
        response = table_income.get_item(Key={'id': income_id})
        income = response.get('Item', {})
    except ClientError as e:
        error_message = f"Error fetching income from DynamoDB: {e.response['Error']['Message']}"
        print(error_message)
        flash('Error fetching income!', 'danger')
        return error_message

    return render_template('edit_income.html', income=income, categories=income_categories)


@app.route('/delete_income/<income_id>', methods=['GET'])
@login_required
def delete_income(income_id):

    table_income = dynamodb.Table(table_income_name)
    table_income.delete_item(Key={'id': income_id})

    flash('Income deleted successfully!', 'success')
    return redirect(url_for('index'))

@app.route('/set_budget', methods=['GET', 'POST'])
@login_required
def set_budget():
    if request.method == 'POST':
        budget_limit = Decimal(request.form['budget_limit'])

        table_budget = dynamodb.Table(table_budget_name)
        table_budget.put_item(
            Item={
                'id': 'monthly_limit',
                'limit': budget_limit
            }
        )

        flash('Budget limit set successfully!', 'success')
        return redirect(url_for('index'))

    return render_template('set_budget.html')

@app.route('/add_expense', methods=['GET', 'POST'])
def add_expense():
    if request.method == 'POST':
        expense_id = request.form['id']
        expense_name = request.form['name']
        expense_amount = Decimal(request.form['amount'])
        expense_category = request.form['category']
        
        table_expenses = dynamodb.Table(table_expenses_name)
        table_expenses.put_item(
            Item={
                'id': expense_id,
                'name': expense_name,
                'amount': expense_amount,
                'category': expense_category
            }
        )
        
      
        expense_data = {
            'id': expense_id,
            'name': expense_name,
            'amount': str(expense_amount),  
            'category': expense_category
        }

       
        send_expense_to_sqs(expense_data)
        
       
        expenses = fetch_all_expenses()  

        budget_limit = 2344
        
        check_budget_limit(expenses, budget_limit)
        
        flash('Expense added successfully!', 'success')
        return redirect(url_for('index'))
    
    return render_template('add_expense.html')

@app.route('/add_income', methods=['GET', 'POST'])
def add_income():
    income_categories = ['Salary', 'Business', 'Investment', 'Other']

    if request.method == 'POST':
        income_id = request.form['id']
        income_name = request.form['name']
        income_amount = Decimal(request.form['amount'])
        income_category = request.form['category'] 

        try:
            table_income = dynamodb.Table(table_income_name)
            table_income.put_item(
                Item={
                    'id': income_id,
                    'name': income_name,
                    'amount': income_amount,  
                    'category': income_category  
                }
            )
            flash('Income added successfully!', 'success')
            return redirect(url_for('index'))
        except ClientError as e:
            error_message = f"Error adding income to DynamoDB: {e.response['Error']['Message']}"
            print(error_message)
            flash('Error adding income!', 'danger')
            return error_message

    return render_template('add_income.html', categories=income_categories)


@app.route('/generate_report', methods=['GET'])
@login_required
def generate_report():
    try:
 
        table_expenses = dynamodb.Table(table_expenses_name)
        table_income = dynamodb.Table(table_income_name)

        expense_response = table_expenses.scan()
        expenses = expense_response.get('Items', [])

        income_response = table_income.scan()
        incomes = income_response.get('Items', [])

        doc = Document()
        doc.add_heading('Financial Report', 0)

        doc.add_heading('Expenses:', level=1)
        for exp in expenses:
            doc.add_paragraph(f"Name: {exp['name']}, Amount: {exp['amount']}, Category: {exp['category']}")

        doc.add_heading('Income:', level=1)
        for inc in incomes:
            doc.add_paragraph(f"Name: {inc['name']}, Amount: {inc['amount']}")

        report_filename = 'financial_report.docx'
        doc.save(report_filename)

        return send_file(report_filename, as_attachment=True)

    except ClientError as e:
        error_message = f"Error generating report: {e.response['Error']['Message']}"
        print(error_message)
        return error_message

if __name__ == '__main__':
    create_dynamodb_tables()
    create_s3_bucket()
    app.run(debug=True)
































